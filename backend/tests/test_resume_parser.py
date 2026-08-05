from docx import Document

from app.resume_parser import PARSER_VERSION, parse_resume


def test_docx_resume_parsing_builds_structured_profile(tmp_path):
    path = tmp_path / "candidate.docx"
    document = Document()
    document.add_heading("Alex Candidate", level=1)
    document.add_paragraph("alex@example.com | +1 555 123 4567")
    document.add_paragraph("Python and FastAPI engineer with 7 years experience using PostgreSQL.")
    document.add_paragraph("Education: Bachelor of Computer Science")
    document.save(path)

    text, profile, version = parse_resume(path, "application/vnd.openxmlformats-officedocument.wordprocessingml.document", ["Python", "FastAPI", "Postgres"], [])

    assert "Alex Candidate" in text
    assert profile["emails"] == ["alex@example.com"]
    assert profile["maximum_stated_years"] == 7
    assert {"python", "fastapi", "postgresql"}.issubset(set(profile["detected_skills"]))
    assert profile["source_format"] == "docx"
    assert version == PARSER_VERSION
