from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / 'docs' / 'SmartHire_Hackathon_Demo_Playbook.docx'
NAVY='082B68'; BLUE='1744BD'; TEAL='07925A'; PALE='F3F7FC'; INK='10204A'; MUTED='59677F'; GOLD='B7791F'

def shade(cell, color):
    tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),color); tcPr.append(shd)
def set_cell_text(cell, text, bold=False, color=INK, size=9.5):
    p=cell.paragraphs[0]; p.paragraph_format.space_after=Pt(0); r=p.add_run(str(text)); r.bold=bold; r.font.name='Calibri'; r._element.rPr.rFonts.set(qn('w:ascii'),'Calibri'); r.font.size=Pt(size); r.font.color.rgb=RGBColor.from_string(color); cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
def borders(table, color='D9E2F0'):
    tblPr=table._tbl.tblPr; b=OxmlElement('w:tblBorders')
    for n in ('top','left','bottom','right','insideH','insideV'):
        e=OxmlElement(f'w:{n}');e.set(qn('w:val'),'single');e.set(qn('w:sz'),'6');e.set(qn('w:color'),color);b.append(e)
    tblPr.append(b)
def table(doc, headers, rows, widths=None):
    t=doc.add_table(rows=1, cols=len(headers));t.alignment=WD_TABLE_ALIGNMENT.CENTER;t.autofit=False;borders(t)
    for i,h in enumerate(headers):
        if widths: t.columns[i].width=Inches(widths[i])
        shade(t.rows[0].cells[i], NAVY);set_cell_text(t.rows[0].cells[i],h,True,'FFFFFF',9)
    for row in rows:
        c=t.add_row().cells
        for i,v in enumerate(row):
            if widths: c[i].width=Inches(widths[i])
            set_cell_text(c[i],v,False,INK,9)
    doc.add_paragraph().paragraph_format.space_after=Pt(3)
    return t
def add_heading(doc,text,level=1):
    p=doc.add_paragraph();p.style=f'Heading {level}';p.paragraph_format.keep_with_next=True;r=p.add_run(text);r.font.color.rgb=RGBColor.from_string(BLUE if level<3 else NAVY);return p
def para(doc,text='',bold_lead=None):
    p=doc.add_paragraph();p.paragraph_format.space_after=Pt(6);p.paragraph_format.line_spacing=1.12
    if bold_lead and text.startswith(bold_lead):
        r=p.add_run(bold_lead);r.bold=True;r.font.color.rgb=RGBColor.from_string(INK);p.add_run(text[len(bold_lead):])
    else:p.add_run(text)
    return p
def bullets(doc, items):
    for x in items:
        p=doc.add_paragraph(style='List Bullet');p.paragraph_format.space_after=Pt(3);p.add_run(x)
def steps(doc, items):
    for x in items:
        p=doc.add_paragraph(style='List Number');p.paragraph_format.space_after=Pt(4);p.add_run(x)
def callout(doc,title,body,color='EAF8F1'):
    t=doc.add_table(rows=1,cols=1);t.alignment=WD_TABLE_ALIGNMENT.CENTER;shade(t.cell(0,0),color);borders(t,'D6E7DD');c=t.cell(0,0);p=c.paragraphs[0];p.paragraph_format.space_after=Pt(2);r=p.add_run(title+' ');r.bold=True;r.font.color.rgb=RGBColor.from_string(NAVY);p.add_run(body);doc.add_paragraph().paragraph_format.space_after=Pt(3)
def image(doc,path,caption):
    p=Path(path)
    if p.exists():
        doc.add_picture(str(p),width=Inches(6.2));q=doc.paragraphs[-1];q.alignment=WD_ALIGN_PARAGRAPH.CENTER
        c=doc.add_paragraph(caption);c.alignment=WD_ALIGN_PARAGRAPH.CENTER;c.runs[0].italic=True;c.runs[0].font.size=Pt(9);c.runs[0].font.color.rgb=RGBColor.from_string(MUTED)

doc=Document();sec=doc.sections[0];sec.top_margin=Inches(.72);sec.bottom_margin=Inches(.72);sec.left_margin=Inches(.78);sec.right_margin=Inches(.78)
styles=doc.styles;styles['Normal'].font.name='Calibri';styles['Normal']._element.rPr.rFonts.set(qn('w:ascii'),'Calibri');styles['Normal'].font.size=Pt(10.5)
for n,size,color in [('Heading 1',16,BLUE),('Heading 2',13,BLUE),('Heading 3',11,NAVY)]:
    s=styles[n];s.font.name='Calibri';s._element.rPr.rFonts.set(qn('w:ascii'),'Calibri');s.font.size=Pt(size);s.font.color.rgb=RGBColor.from_string(color)
header=sec.header.paragraphs[0];header.text='SMARTHIRE  |  HACKATHON DEMO PLAYBOOK';header.runs[0].font.size=Pt(8);header.runs[0].font.color.rgb=RGBColor.from_string(MUTED)
footer=sec.footer.paragraphs[0];footer.alignment=WD_ALIGN_PARAGRAPH.CENTER;footer.add_run('SmartHire - AI-assisted recruitment decision support')

p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;p.paragraph_format.space_before=Pt(76);r=p.add_run('SMARTHIRE');r.bold=True;r.font.size=Pt(15);r.font.color.rgb=RGBColor.from_string(TEAL)
p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;r=p.add_run('Hackathon Demo Playbook');r.bold=True;r.font.size=Pt(30);r.font.color.rgb=RGBColor.from_string(NAVY)
p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;r=p.add_run('A complete presenter guide for the AI-powered recruitment platform');r.font.size=Pt(14);r.font.color.rgb=RGBColor.from_string(MUTED)
doc.add_paragraph();table(doc,['Demo objective','Audience','Suggested duration'],[['Demonstrate an explainable, multi-tenant recruitment product from sign-up to admin governance.','Hackathon judges, recruiters, technical reviewers, and potential users.','12-15 minutes live demo + 3-5 minutes Q&A.']],[2.3,2.3,1.9])
callout(doc,'Presenter promise:','SmartHire does not automate hiring or rejection. It gives people structured, explainable decision support and keeps final decisions with recruiters.','FFF8E8')
doc.add_page_break()

add_heading(doc,'1. Product at a glance');para(doc,'SmartHire is a role-based recruitment platform built for applicants, employers, and platform administrators. It combines job discovery, candidate management, explainable resume matching, workspace controls, subscription governance, and auditability in one product.')
table(doc,['Role','Primary value','What to demonstrate'],[['Applicant','Find relevant work and understand applications.','Profile, job board, salary, resume upload, match report, notifications.'],['Employer','Run a structured hiring process.','Workspace, jobs, ATS pipeline, scorecards, candidates, plan requests.'],['Administrator','Protect platform quality and commercial governance.','Admin setup, user controls, operations, SaaS plans, approvals, auditability.']],[1.1,2.7,2.7])
add_heading(doc,'2. Demo setup before recording');steps(doc,['Start the backend and frontend. Confirm /health returns healthy and open the application in a fresh browser session.','Prepare one Employer account, one Applicant account, and one Admin account. Use realistic but safe demo data.','Create at least one published job: Senior Backend Engineer, salary PKR 150,000-250,000/month, skills Python, FastAPI, PostgreSQL.','Prepare a text-based PDF resume containing some, but not all, required skills.','Prepare one pending plan-change request so the approval flow can be shown without waiting.','Open the app in a private/incognito browser for applicant sign-up if you want to demonstrate separate sessions.'])
callout(doc,'Recording tip:','Keep three tabs open: Employer workspace, Applicant workspace, and Admin Control Center. This makes role switching fast and prevents accidental logout.','EAF3FF')
add_heading(doc,'3. Suggested live-demo story');para(doc,'Use a simple story: “A growing university or technology company needs a backend engineer. The employer creates a transparent job posting. An applicant applies. SmartHire analyzes the resume with evidence. The recruiter manages the candidate. The employer requests a plan upgrade. The platform admin reviews and approves it.”')

add_heading(doc,'4. Employer flow');add_heading(doc,'4.1 Employer account creation and verification',2);steps(doc,['Open Register and select Employer.','Enter name, work email, company name, and password; alternatively show Google registration.','Verify the email with the code or verification link.','Explain that an employer workspace is created and the owner receives workspace management permissions.'])
add_heading(doc,'4.2 Employer dashboard',2);bullets(doc,['Active jobs, total applicants, average match score, and total jobs provide the opening operational summary.','Recent jobs provide access to candidate lists.','The left navigation exposes My Jobs, Post a Job, Hiring Pipeline, Recruiter Team, AI Scorecards, Workspace & Billing, Notifications, and Profile.'])
add_heading(doc,'4.3 Post a transparent job',2);steps(doc,['Open Post a Job.','Enter the title, detailed description, required skills, location, employment type, experience level, and status.','Enter minimum and maximum monthly salary in PKR. Explain that salary is shown to applicants before they apply.','Publish the job. Then open My Jobs and show the salary column and applicant count.'])
image(doc,r'C:\Users\USER\Pictures\Screenshots\Screenshot 2026-08-08 073301.png','Figure 1. Employer job-posting experience. Salary fields are included in the current product flow.')
add_heading(doc,'4.4 Employer ATS and AI scorecard',2);bullets(doc,['Hiring Pipeline supports recruiter workflow stages, such as Applied, Screening, Interview, Offer, Hired, and Rejected.','Candidate workspace supports assignments, tags, private notes, and an activity timeline.','AI Scorecards let the employer configure relevance, skills, experience, role alignment, domain knowledge, and education/certification weights.','Explain that score changes do not make automatic hiring decisions; human review remains required.'])

add_heading(doc,'5. Applicant flow');add_heading(doc,'5.1 Applicant registration and profile',2);steps(doc,['Open Register and choose Applicant.','Create the account or use Google registration.','Verify the email.','Open Profile and show the applicant fields: contact information, headline, skills, education, languages, experience, availability, work preferences, and portfolio links.'])
add_heading(doc,'5.2 Browse, inspect, and apply',2);steps(doc,['Open Job Board and search for the backend role.','Open job details. Highlight company, location, employment type, skills, experience level, and salary range.','Upload the prepared PDF/DOCX resume and submit.','Explain the duplicate-application protection: the applicant cannot apply twice to the same job.'])
add_heading(doc,'5.3 Applicant results and transparency',2);bullets(doc,['My Applications shows processing state, final match score, and AI insight.','The job details page keeps salary visible at application time.','The explainable report shows component scores, matched skills, evidence, missing evidence, strengths, gaps, provider state, and recommendation.','If an application is already submitted, the user sees a clear status rather than another upload form.'])

add_heading(doc,'6. Explainable AI matching');para(doc,'Every valid resume is evaluated by a deterministic evidence engine. Gemini enrichment is optional and never replaces the deterministic result.')
table(doc,['Deterministic component','Default weight','How it is measured'],[['Skills evidence','35%','Explicit skill evidence, alias normalization, and mandatory/preferred/optional priorities.'],['Experience','25%','Explicitly stated years compared with role seniority.'],['Semantic relevance','15%','TF-IDF cosine similarity: 70% word n-grams and 30% character n-grams.'],['Role alignment','15%','Meaningful job-title term coverage in the resume.'],['Domain knowledge','5%','Employer-configured domain keyword evidence.'],['Education/certifications','5%','Configured requirement coverage.']],[2.2,1.0,3.3])
para(doc,'Formula: Deterministic score = sum(component score x component weight). A missing mandatory skill caps the deterministic score at 55%.')
table(doc,['Hybrid guardrail','Current behavior'],[['Maximum Gemini influence','35% before confidence adjustment.'],['Effective AI influence','Maximum AI weight x Gemini confidence.'],['Manual review trigger','Difference of 25 points or more between deterministic and AI scores.'],['AI failure behavior','Safe deterministic fallback; the application still completes.'],['Protected traits','Name, email, age, gender, nationality, and other protected traits are excluded from scoring.']],[2.4,4.1])
callout(doc,'Judge-ready explanation:','“The score is not a hiring decision. It is an auditable evidence summary. Recruiters can see why it scored that way and must retain human judgment.”','FFF8E8')

add_heading(doc,'7. Admin Control Center');add_heading(doc,'7.1 Secure first-admin creation',2);steps(doc,['Open /admin/setup only when the database has no administrator.','Enter name, email, strong password, and the private ADMIN_BOOTSTRAP_TOKEN stored in environment configuration.','Create the first administrator, then sign in at /admin/login.','Explain that public registration never creates administrators and the setup route closes once an admin exists.'])
image(doc,r'C:\Users\USER\Pictures\Screenshots\Screenshot 2026-08-08 074530.png','Figure 2. Separate Admin Control Center login experience.')
add_heading(doc,'7.2 Admin capabilities and protections',2);table(doc,['Area','Admin capability','Protection'],[['Identity & Access','View users; activate or suspend Applicants and Employers; create additional admins.','Admins cannot change an Employer/Applicant into an Admin. Admin accounts, including the current admin, cannot be suspended, activated, or role-changed through user management.'],['Content Control','Moderate platform job content and review job/application state.','Administrative RBAC is enforced by backend endpoints.'],['Operations & Audit','Review activity, job processing, queue state, and platform health.','Audit records capture sensitive actions.'],['Intelligence Monitor','Review score distribution, AI completion, disagreement, manual-review volume, and overrides.','Protected characteristics are intentionally excluded.'],['SaaS Accounts','View organizations, plans, status, plan mix, and estimated MRR.','Commercial changes require approval workflow.']],[1.25,2.65,2.6])

add_heading(doc,'8. SaaS plan governance demo');steps(doc,['As Employer, open Workspace & Billing. Show Starter, Growth, and Scale plan cards, pricing, entitlement limits, current plan, and usage.','Select a different plan. Explain that this now creates a request rather than activating immediately.','Show the immediate browser confirmation and the employer Notification Center item.','As Admin, open the dedicated Plan Requests sidebar item. Show workspace, requester, current plan, requested plan, timestamp, and status.','Approve or reject the request. On approval, explain that the requested plan activates and the employer receives an in-app notification.','Return to Employer Workspace & Billing and show the updated plan only after approval.'])
table(doc,['Plan','Price','Entitlements'],[['Starter','$0/month','3 active jobs; 3 team members; 50 AI analyses/month; 500 MB storage.'],['Growth','$99/month','25 active jobs; 15 team members; 1,500 AI analyses/month; 10 GB storage.'],['Scale','$299/month','Unlimited active jobs; unlimited team members; unlimited AI analyses; 100 GB storage.']],[1.0,1.0,4.5])
image(doc,r'C:\Users\USER\Pictures\Screenshots\Screenshot 2026-08-08 082908.png','Figure 3. Dedicated Plan Requests navigation in the Admin Control Center.')

add_heading(doc,'9. Authentication, privacy, and notifications');bullets(doc,['Password sign-in uses hashed passwords and JWT access/refresh sessions.','Email verification is required for normal password registration.','Google OAuth supports Applicant and Employer registration/login; administrators use the separate Control Center login.','Forgot Password creates one-time, expiring reset links and revokes existing refresh sessions after a successful password reset.','Notifications are stored in-app. Plan requests notify the employer and all active administrators; approval/rejection notifies the requesting employer.','Resumes are private; employers access only authorized candidate data within their workspace.'])

add_heading(doc,'10. Recording script and timing');table(doc,['Time','Screen / action','Presenter line'],[['0:00-0:45','Landing + role overview','“SmartHire supports the full hiring journey: applicants, employers, and governed platform administration.”'],['0:45-2:30','Employer registration + dashboard','“A verified employer gets an isolated workspace and structured hiring tools.”'],['2:30-4:00','Post job with salary','“Transparency begins before application: salary, requirements, and experience expectations are visible.”'],['4:00-6:00','Applicant registration + application','“Applicants discover roles, inspect salary and requirements, then submit a private resume.”'],['6:00-8:30','Match report','“This is explainable decision support: evidence, weights, missing skills, and human-review guardrails.”'],['8:30-10:00','Employer ATS pipeline','“Recruiters move from scored application to collaborative candidate workflow.”'],['10:00-12:00','Employer plan request + Admin approval','“Commercial actions are governed: employers request; admins review; only approval activates.”'],['12:00-13:30','Admin controls + close','“The Control Center protects identities, content, operations, intelligence, and subscriptions.”']],[.7,2.0,3.8])

add_heading(doc,'11. Hackathon Q&A preparation');table(doc,['Likely question','Short answer'],[['Does AI make the hiring decision?','No. SmartHire is decision support. Recruiters remain responsible for hiring decisions, and score disagreement triggers manual review.'],['How is the score explainable?','The platform stores weighted components, evidence snippets, skill coverage, scorecard snapshot, AI confidence, and hybrid weights.'],['How is tenant data isolated?','Jobs, candidates, pipelines, memberships, notifications, and plan data are organization-scoped; permissions are checked server-side.'],['How do you prevent unauthorized admin access?','No public admin registration. The first admin requires a private bootstrap token; later admins are created only by authenticated admins.'],['How does monetization work?','Employer workspaces have Starter, Growth, and Scale entitlements. Plan changes require admin approval in the current controlled workflow.'],['What is next after the hackathon?','Persistent private object storage, real payment-provider checkout, production email, managed workers, observability, and evaluation against human-reviewed hiring outcomes.']],[2.35,4.15])

add_heading(doc,'12. Final pre-demo checklist');bullets(doc,['Backend health endpoint is healthy and frontend is open in a clean browser.','All demo accounts are verified and passwords are known.','A job with salary is published.','A resume is ready and text-based.','At least one applicant application has completed analysis.','At least one plan request is pending or ready to create.','Admin Control Center is available and Plan Requests is visible in the sidebar.','No API keys, tokens, passwords, or database credentials are visible in the recording.','Use the demo story, not every feature: clarity is stronger than feature volume.'])
callout(doc,'Closing line:','“SmartHire brings transparent job discovery, explainable AI assistance, collaborative ATS workflows, and platform governance into one recruitment experience.”','EAF8F1')
doc.save(OUT)
print(OUT)
